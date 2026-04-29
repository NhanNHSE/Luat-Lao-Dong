"""Document processing service: extract text from PDF, DOCX, images with OCR."""

import io
import base64
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
from google import genai
from google.genai.types import GenerateContentConfig, Part

from src.core.config import get_settings
from src.core.logging import get_logger

settings = get_settings()
logger = get_logger(__name__)

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".doc": "application/msword",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".tiff": "image/tiff",
    ".tif": "image/tiff",
}

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB

_client = None


def _get_client() -> genai.Client:
    """Get or create the Gemini client."""
    global _client
    if _client is None:
        _client = genai.Client(api_key=settings.gemini_api_key)
    return _client


def validate_file(filename: str, file_size: int) -> Optional[str]:
    """Validate file type and size.

    Args:
        filename: Original filename.
        file_size: File size in bytes.

    Returns:
        Error message string if invalid, None if valid.
    """
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(SUPPORTED_EXTENSIONS.keys())
        return f"Định dạng file không được hỗ trợ. Chấp nhận: {supported}"

    if file_size > MAX_FILE_SIZE:
        return f"File quá lớn. Giới hạn: {MAX_FILE_SIZE // (1024*1024)}MB"

    return None


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF. Falls back to OCR for scanned pages.

    Args:
        file_bytes: Raw PDF file bytes.

    Returns:
        Extracted text content.
    """
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    ocr_pages = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text().strip()

        if text and len(text) > 50:
            # Page has extractable text
            text_parts.append(f"--- Trang {page_num + 1} ---\n{text}")
        else:
            # Page is likely scanned/image - need OCR
            ocr_pages.append(page_num)

    # OCR scanned pages using Gemini Vision
    if ocr_pages:
        logger.info("ocr_needed", pages=ocr_pages, total_pages=len(doc))
        for page_num in ocr_pages:
            page = doc[page_num]
            # Render page to image at 200 DPI
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")

            ocr_text = ocr_image_with_gemini(img_bytes, "image/png")
            if ocr_text:
                text_parts.append(f"--- Trang {page_num + 1} (OCR) ---\n{ocr_text}")

    doc.close()

    if not text_parts:
        return ""

    return "\n\n".join(text_parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file.

    Args:
        file_bytes: Raw DOCX file bytes.

    Returns:
        Extracted text content.
    """
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def extract_text_from_doc(file_bytes: bytes) -> str:
    """Extract text from old .doc (binary) format using antiword.

    Falls back to reading raw text if antiword is not available.

    Args:
        file_bytes: Raw DOC file bytes.

    Returns:
        Extracted text content.
    """
    # Try antiword first
    try:
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        result = subprocess.run(
            ['antiword', tmp_path],
            capture_output=True, text=True, timeout=30
        )
        import os
        os.unlink(tmp_path)

        if result.returncode == 0 and result.stdout.strip():
            logger.info("doc_extracted_antiword", chars=len(result.stdout))
            return result.stdout.strip()
    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
        logger.warning("antiword_unavailable", error=str(e))

    # Fallback: try python-docx anyway (works for some .doc files saved as .docx)
    try:
        return extract_text_from_docx(file_bytes)
    except Exception:
        pass

    # Last resort: extract readable text from binary
    try:
        text = file_bytes.decode('utf-8', errors='ignore')
        # Filter printable lines
        lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 10]
        if lines:
            return "\n".join(lines)
    except Exception:
        pass

    raise ValueError("Không thể đọc file .doc. Vui lòng chuyển sang .docx hoặc .pdf.")


def extract_text_from_image(file_bytes: bytes, mime_type: str) -> str:
    """Extract text from image using Gemini Vision OCR.

    Args:
        file_bytes: Raw image file bytes.
        mime_type: MIME type of the image.

    Returns:
        Extracted text content.
    """
    return ocr_image_with_gemini(file_bytes, mime_type)


def ocr_image_with_gemini(image_bytes: bytes, mime_type: str) -> str:
    """Use Gemini Vision to OCR text from an image.

    Args:
        image_bytes: Raw image bytes.
        mime_type: Image MIME type.

    Returns:
        Extracted text from the image.
    """
    client = _get_client()

    prompt = """Hãy đọc và trích xuất TOÀN BỘ nội dung văn bản trong hình ảnh này.
Giữ nguyên cấu trúc văn bản gốc (tiêu đề, đoạn, danh sách, bảng).
Nếu có bảng, hãy format dạng markdown table.
Chỉ trả về nội dung văn bản, không thêm bình luận hay giải thích."""

    image_part = Part.from_bytes(data=image_bytes, mime_type=mime_type)

    try:
        response = client.models.generate_content(
            model=settings.llm_model,
            contents=[prompt, image_part],
            config=GenerateContentConfig(
                temperature=0.1,
                max_output_tokens=4096,
            ),
        )
        return response.text.strip()
    except Exception as e:
        logger.error("ocr_failed", error=str(e))
        return ""


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Extract text from a file based on its extension.

    Args:
        filename: Original filename.
        file_bytes: Raw file bytes.

    Returns:
        Extracted text content.
    """
    ext = Path(filename).suffix.lower()
    logger.info("extracting_text", filename=filename, ext=ext, size=len(file_bytes))

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    elif ext == ".doc":
        return extract_text_from_doc(file_bytes)
    elif ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".tif"):
        mime_type = SUPPORTED_EXTENSIONS.get(ext, "image/png")
        return extract_text_from_image(file_bytes, mime_type)
    else:
        return ""
